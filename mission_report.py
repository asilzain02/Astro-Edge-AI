import os
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

class MissionReport:
    def __init__(self, save_dir="reports"):
        self.save_dir = save_dir
        os.makedirs(self.save_dir, exist_ok=True)

    def _filename(self, ext):
        date_str = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        return os.path.join(self.save_dir, f"D:\astro_edge_ai\astro_edge_ai\logs\mission_report_{date_str}.{ext}")

    def export_pdf(self, logs):
        filename = self._filename("pdf")
        doc = SimpleDocTemplate(filename, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph("<b>Mission Report</b>", styles["Title"]))
        story.append(Spacer(1, 12))

        for log in logs:
            story.append(Paragraph(f"<b>User:</b> {log['query']}", styles["Normal"]))
            story.append(Paragraph(f"<b>Assistant:</b> {log['response']}", styles["Normal"]))
            if "objects" in log and log["objects"]:
                story.append(Paragraph(f"<b>Detected Objects:</b> {', '.join(log['objects'])}", styles["Normal"]))

            if "metrics" in log:
                metrics_str = ", ".join([f"{k}: {v}" for k, v in log["metrics"].items()])
                story.append(Paragraph(f"<b>System Metrics:</b> {metrics_str}", styles["Normal"]))

        doc.build(story)
        return filename

    def export_word(self, logs):
        filename = self._filename("docx")
        doc = Document()
        doc.add_heading("Mission Report", 0)

        for log in logs:
            doc.add_heading("User Query", level=2)
            doc.add_paragraph(log["query"])
            doc.add_heading("Assistant Response", level=2)
            doc.add_paragraph(log["response"])
            # Inside export_word
            if "objects" in log and log["objects"]:
                doc.add_heading("Detected Objects", level=2)
                doc.add_paragraph(", ".join(log["objects"]))

            if "metrics" in log:
                doc.add_heading("System Metrics", level=2)
                for k, v in log["metrics"].items():
                    doc.add_paragraph(f"{k}: {v}")


        doc.save(filename)
        return filename
